import io
import os
import docx
from PyPDF2 import PdfReader
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
import streamlit as st
import pandas as pd
from fpdf import FPDF

# ---------- Setup ----------
model = SentenceTransformer('all-MiniLM-L6-v2')

CONSULTANT_FOLDER = "./Consultant_Profiles"
CLIENT_FOLDER = "./Client_JD"

os.makedirs(CONSULTANT_FOLDER, exist_ok=True)
os.makedirs(CLIENT_FOLDER, exist_ok=True)

def read_pdf(path_or_file):
    reader = PdfReader(path_or_file)
    text = ""
    for page in reader.pages:
        if page.extract_text():
            text += page.extract_text()
    return text

def read_docx(path_or_file):
    doc = docx.Document(path_or_file)
    return "\n".join([p.text for p in doc.paragraphs])

def make_summary(text, max_chars=200):
    text = text.strip().replace("\n", " ")
    return text[:max_chars] + ("..." if len(text) > max_chars else "")

def get_embedding(text):
    return model.encode(text)

def generate_jd(requirement_text):
    return f"""
Title: Data Analyst
Location: (unspecified)
Experience: Minimum 3 years
Responsibilities:
- Analyze data and prepare reports
- Support decision-making with insights
- Collaborate with teams on projects
Qualifications:
- Strong analytical skills
- Proficiency in data tools
- Excellent communication
"""

def jd_to_pdf(jd_text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, jd_text)
    pdf_bytes = pdf.output(dest='S').encode('latin1')
    return pdf_bytes

def jd_to_docx(jd_text):
    doc = docx.Document()
    doc.add_paragraph(jd_text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def load_consultant_folder():
    consultant_data = []
    if os.path.exists(CONSULTANT_FOLDER):
        for fname in os.listdir(CONSULTANT_FOLDER):
            path = os.path.join(CONSULTANT_FOLDER, fname)
            if fname.endswith(".pdf"):
                text = read_pdf(path)
            elif fname.endswith(".docx"):
                text = read_docx(path)
            else:
                continue
            consultant_data.append((fname, get_embedding(text), make_summary(text)))
    return consultant_data

def load_client_folder():
    client_data = []
    if os.path.exists(CLIENT_FOLDER):
        for fname in os.listdir(CLIENT_FOLDER):
            path = os.path.join(CLIENT_FOLDER, fname)
            if fname.endswith(".pdf"):
                text = read_pdf(path)
            elif fname.endswith(".docx"):
                text = read_docx(path)
            else:
                continue
            client_data.append((fname, get_embedding(text), make_summary(text)))
    return client_data

# ---------- Streamlit UI ----------
st.title("Consultant Matching Demo")

# Mode toggle
mode = st.radio("Select Mode", ["Client", "Consultant"])

# ---------- Client Mode ----------
if mode == "Client":
    client_files = st.file_uploader("Upload Client JDs", type=["pdf","docx"], accept_multiple_files=True)
    client_text = st.text_area("Or type your requirements here (free text)", "")
    consultant_data = load_consultant_folder()
    client_data = []   # start empty, only fill from uploads or generated JD

    if st.button("Run Matching"):
        if client_files:
            for file in client_files:
                save_path = os.path.join(CLIENT_FOLDER, file.name)
                with open(save_path, "wb") as f:
                    f.write(file.getbuffer())
                if file.name.endswith(".pdf"):
                    text = read_pdf(save_path)
                else:
                    text = read_docx(save_path)
                client_data.append((file.name, get_embedding(text), make_summary(text)))

        if client_text.strip():
            jd_text = generate_jd(client_text)
            st.subheader("Generated Job Description from Requirement")
            st.text(jd_text)

            pdf_bytes = jd_to_pdf(jd_text)
            st.download_button("Download JD as PDF", pdf_bytes, "generated_jd.pdf", "application/pdf")

            docx_bytes = jd_to_docx(jd_text)
            st.download_button("Download JD as Word", docx_bytes, "generated_jd.docx",
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            client_data.append(("Generated JD", get_embedding(jd_text), make_summary(jd_text)))

        if not consultant_data:
            st.warning("⚠️ No consultant profiles found.")
        if not client_data:
            st.warning("⚠️ No client JDs found.")

        if consultant_data and client_data:
            dimension = len(consultant_data[0][1])
            index_consultants = faiss.IndexFlatL2(dimension)
            consultant_embeddings = np.array([c[1] for c in consultant_data])
            index_consultants.add(consultant_embeddings)

            results_consultants = []
            for client_file, client_emb, client_summary in client_data:
                D, I = index_consultants.search(np.array([client_emb]), 3)
                if len(I[0]) == 0:
                    st.warning(f"⚠️ No consultant match found for {client_file}.")
                else:
                    consultant_names = [consultant_data[idx][0] for idx in I[0]]
                    scores = [100 - dist for dist in D[0]]
                    best_idx = np.argmax(scores)
                    st.success(f"⭐ Best Consultant Match for {client_file}: {consultant_names[best_idx]} ({scores[best_idx]:.2f}%)")

                    for dist, idx in zip(D[0], I[0]):
                        similarity = 100 - dist
                        consultant_file, _, consultant_summary = consultant_data[idx]
                        results_consultants.append([client_file, client_summary, consultant_file, consultant_summary, f"{similarity:.2f}%"])

            if results_consultants:
                st.subheader("Client → Consultant Matches")
                df_cons = pd.DataFrame(results_consultants, columns=["Client JD", "Client Summary", "Consultant Profile", "Consultant Summary", "Match Strength"])
                st.dataframe(df_cons)
                csv = df_cons.to_csv(index=False).encode("utf-8")
                st.download_button("Download JD→Consultant Results as CSV", csv, "jd_to_consultant_results.csv", "text/csv")

# ---------- Consultant Mode ----------
elif mode == "Consultant":
    consultant_files = st.file_uploader("Upload Consultant Profiles (CVs)", type=["pdf","docx"], accept_multiple_files=True)
    client_data = load_client_folder()
    consultant_data = []   # start empty, only fill from uploads

    if st.button("Run Matching"):
        if consultant_files:
            for file in consultant_files:
                save_path = os.path.join(CONSULTANT_FOLDER, file.name)
                with open(save_path, "wb") as f:
                    f.write(file.getbuffer())
                if file.name.endswith(".pdf"):
                    text = read_pdf(save_path)
                else:
                    text = read_docx(save_path)
                consultant_data.append((file.name, get_embedding(text), make_summary(text)))

        if not client_data:
            st.warning("⚠️ No client JDs found.")
        if not consultant_data:
            st.warning("⚠️ No consultant profiles found.")

        if consultant_data and client_data:
            dimension2 = len(client_data[0][1])
            index_clients = faiss.IndexFlatL2(dimension2)
            client_embeddings = np.array([c[1] for c in client_data])
            index_clients.add(client_embeddings)

            results_clients = []
            for consultant_file, consultant_emb, consultant_summary in consultant_data:
                D2, I2 = index_clients.search(np.array([consultant_emb]), 3)
                if len(I2[0]) == 0:
                    st.warning(f"⚠️ No JD match found for {consultant_file}.")
                else:
                    jd_names = [client_data[idx][0] for idx in I2[0]]
                    scores2 = [100 - dist for dist in D2[0]]
                    best_idx2 = np.argmax(scores2)
                    st.success(f"📄 Best JD Match for {consultant_file}: {jd_names[best_idx2]} ({scores2[best_idx2]:.2f}%)")

                    for dist, idx in zip(D2[0], I2[0]):
                        similarity = 100 - dist
                        jd_file, _, jd_summary = client_data[idx]
                        results_clients.append([consultant_file, consultant_summary, jd_file, jd_summary, f"{similarity:.2f}%"])

            if results_clients:
                st.subheader("Consultant → JD Matches")
                df_cli = pd.DataFrame(results_clients, columns=["Consultant Profile", "Consultant Summary", "Client JD", "Client Summary", "Match Strength"])
                st.dataframe(df_cli)
                csv2 = df_cli.to_csv(index=False).encode("utf-8")
                st.download_button("Download Consultant→JD Results as CSV", csv2, "consultant_to_jd_results.csv", "text/csv")
                